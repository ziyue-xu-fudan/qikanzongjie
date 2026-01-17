from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_data_request_letter():
    doc = Document()
    
    # 设置字体样式
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # 1. 头部信息 (Sender Info) - 右对齐
    header_info = [
        "Zhi-Ming Shao, M.D., Ph.D.",
        "Director, Department of Breast Surgery",
        "Fudan University Shanghai Cancer Center",
        "No. 270 Dong’an Rd, Shanghai 200032, China.",
        "Phone: 8621-6417-5590;",
        "Fax: 8621-6443-4556;",
        "E-mail: zhimin_shao@yeah.net;"
    ]
    
    for line in header_info:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)

    doc.add_paragraph()  # 空行

    # 2. 收件人信息
    recipient_info = [
        "To the Data Access Committee / Corresponding Author,",
        "Authors of 'Whole-genome landscapes of 1,364 breast cancers'",
    ]
    
    for line in recipient_info:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(0)
    
    doc.add_paragraph() # 空行

    # 3. 称呼
    doc.add_paragraph("Dear Data Access Committee and Authors,")

    # 4. 正文内容
    
    # Paragraph 1: Purpose
    p1 = doc.add_paragraph(
        "I am writing to formally request access to the genomic and clinical data associated with your recently published article: "
        "\"Whole-genome landscapes of 1,364 breast cancers\" (Nature, 2024). "
        "We have read your work with great interest and believe it represents a landmark contribution to understanding the genomic heterogeneity of breast cancer."
    )
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Paragraph 2: Importance of their data (Multi-ethnic focus)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    runner = p2.add_run("Your study provides a unique and comprehensive resource for the ")
    runner = p2.add_run("genomic landscape of Asian breast cancer populations")
    runner.bold = True
    runner = p2.add_run(", which has been historically underrepresented in major Western cohorts. ")
    runner = p2.add_run("Access to this high-quality, large-scale whole-genome sequencing (WGS) dataset is critical for establishing a truly multi-center, multi-ethnic perspective on breast cancer genomics.")
    runner.bold = True

    # Paragraph 3: Our Background & Specific Aim
    p3 = doc.add_paragraph(
        "Our research group at Fudan University Shanghai Cancer Center has established one of the largest multi-omic atlases of breast cancer in China (Cancer Cell, 2019; Nat Genetics, 2023) and has successfully translated these findings into subtype-based treatment strategies (Lancet Oncol, 2023; JAMA, 2024). "
    )
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    runner = p3.add_run("We aim to integrate your dataset with our in-house multi-omics cohorts to validate population-specific driver mutations and explore the cross-ethnic heterogeneity of structural variants.")
    runner.bold = True
    runner = p3.add_run(" This comparative analysis will enable us to identify robust therapeutic targets that are applicable across diverse populations.")

    # Paragraph 4: Collaboration & Compliance
    p4 = doc.add_paragraph(
        "We strictly adhere to data protection regulations and will ensure that all data is used solely for non-commercial, academic research purposes. We will ensure full attribution and citation of your work in any resulting publications. "
    )
    p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    runner = p4.add_run("Furthermore, we are highly open to potential collaboration opportunities and would be delighted to share our findings and discuss future joint efforts to advance precision oncology for diverse patient populations.")
    runner.bold = True

    p5 = doc.add_paragraph(
        "We are happy to provide any further documentation required for the data access agreement. Thank you for your time and consideration."
    )
    
    doc.add_paragraph() # 空行

    # 5. 签名
    doc.add_paragraph("Sincerely,")
    
    # 签名块
    signature_lines = [
        "Zhi-Ming Shao, M.D., Ph.D.",
        "Director, Department of Breast Surgery",
        "Fudan University Shanghai Cancer Center and Cancer Institute",
        "Professor of Oncology, Shanghai Medical College, Fudan University"
    ]
    
    for line in signature_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(0)

    # 保存文件到桌面
    desktop_path = os.path.expanduser("~/Desktop")
    file_name = "Data_Access_Request_Letter_Revised.docx"
    save_path = os.path.join(desktop_path, file_name)
    
    try:
        doc.save(save_path)
        print(f"Document saved successfully to: {save_path}")
    except Exception as e:
        print(f"Error saving document: {e}")
        fallback_path = os.path.join(os.getcwd(), file_name)
        doc.save(fallback_path)
        print(f"Document saved to fallback path: {fallback_path}")

if __name__ == "__main__":
    create_data_request_letter()
